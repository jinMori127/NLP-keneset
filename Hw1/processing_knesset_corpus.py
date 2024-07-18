import zipfile
import os
from docx import Document
import re
import pandas as pd
import sys

punctuation_marks = ['. ', '! ', '? ']
not_word = ['.', '!', '?', '"', ';']
list_pre = ["א'", "ב'", "ג'", "ד'", "ה'", "ו'", "ז'", "ח'", "ט'", "י'", "כ'", "ל'", "מ'", "נ'", "ס'", "ע'", "ש'", "ק'",
            "ר'", "פ'", "ת'"]


# input: text (speaker text)
# output: list of sentences
def split_sentence(text):
    sentences_list = []

    # split the text using the specified punctuation marks and check for double quotes
    text_split = re.split(r'(?<=[.!?])\s*(?=\")|(?<=[.!?])\s+', text)

    # remove redundant whitespaces and double quotes
    text_split = [sentence.strip(' "') for sentence in text_split]
    sentences_list.extend(text_split)

    return sentences_list


# takes as an inout a sentence return a list of tokens
def tokenized(text):
    # tokenize based on whitespace and punctuation
    tokens = re.findall(r'\b\w+\b|[.,;!?"]', text, re.UNICODE)

    last_indx = 0
    for index in range(len(text)):
        if '"' == text[index]:
            # check if '"' comes after a Hebrew character, if yes, keep them together; else, split
            if last_indx == 0:
                last_indx = tokens.index('"')

            else:
                if last_indx + 1 < len(tokens) and '"' in tokens[last_indx + 1:]:
                    last_indx = last_indx + tokens[last_indx + 1:].index('"') + 1
            if last_indx + 2 < len(tokens):
                if (index + 1 < len(text)) and ('\u0590' <= text[index + 1] <= '\u05FF') and (
                        text[index - 1] != " "):
                    inx = last_indx
                    tokens = tokens[:inx - 1] + [tokens[inx - 1] + tokens[inx] + tokens[inx + 1]] + tokens[inx + 2:]
    return tokens


# make clean for the sentence
# if the sentence contain english words it will delete the sentence retval ""
# else will clean -- and the redundant spaces  retval "cleaned_text"
def clean_sentence(text):
    has_english = re.search(r'[a-zA-Z]', text) is not None
    if has_english:
        return ""

    allowed_characters = re.sub(r'[^\u0590-\u05FF\d\s.,،?!":%-]', '', text)

    cleaned_text_1 = allowed_characters.replace('--', ' ')
    cleaned_text_2 = re.sub(r'\s+\.', '.', cleaned_text_1)
    # Replace consecutive spaces with a single space
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text_2)

    return cleaned_text


# input list of tokens
# output the number of words
def calculate_sentence_length(token_list):
    count = 0
    for word in token_list:
        if word not in not_word:
            count += 1
    return count


# takes a file and protocol type and return a map with the names and sentences
def extract_names_and_sentences(document, protocol_type):
    # this map going to be used to save the names as a keys and
    # the spoken text as value of the key
    speakers_map = {}
    last_name = ""
    text_before_colon = ""
    is_changed = False
    # extracting the names and the text for each speaker
    if protocol_type == "plenary":
        for par in document.paragraphs:
            # check if the text is not Heading or Title or Normal format and have underline
            # because in this format the speaker names come in the ptm files
            if par.style and not (
                    par.style.name.startswith("Heading") or par.style.name == "Title" or par.style.name == "Normal"):
                if ':' in par.text:
                    is_changed = True
                    for prefix in list_pre:
                        if prefix in par.text:
                            # Modify the text_before_colon line to remove everything before the prefix
                            modified_sentence = par.text.split(':')[0].split(prefix)[-1].split('(', 1)[0].strip()
                            text_before_colon = re.sub(r'\b' + re.escape('היו”ר') + r'\b', '', modified_sentence)
                            last_name = prefix + " " + text_before_colon
                            break  # break out of the loop once a prefix is found

                    if text_before_colon == "":
                        # If "ד' ש'" is not present, use the entire text before colon
                        text_before_colon = par.text.split(':')[0].strip().split('(', 1)[0].strip()
                        text_before_colon = re.sub(r'\b' + re.escape('היו”ר') + r'\b', '', text_before_colon)

                        last_name = text_before_colon
                    # adding the name for the mao list
                    if last_name not in speakers_map.keys():
                        speakers_map[last_name] = ""

                    text_before_colon = ""
            if par.style and par.style.name == "Normal" and last_name != "":
                speakers_map[last_name] += par.text
                if is_changed:
                    speakers_map[last_name] += " "
    else:
        for par in document.paragraphs:
            # check if the text is not in bold format and have underline
            # because in this format the speaker names come in the ptv files
            if (not any(run.bold for run in par.runs)) and any(run.underline for run in par.runs):
                if ':' in par.text:
                    is_changed = True

                    modified_sentence = re.sub(r'\b' + re.escape('היו”ר') + r'\b', '', par.text)
                    modified_name = modified_sentence.lstrip('1')
                    modified_name = re.sub(r'\b' + re.escape('היו"ר') + r'\b', '', modified_name)
                    text_before_colon = modified_name.split(':')[0].strip()
                if text_before_colon not in speakers_map.keys():
                    speakers_map[text_before_colon] = ""

            elif par.style and par.style.name == "Normal" and text_before_colon != "":
                speakers_map[text_before_colon] += par.text.strip()
                if is_changed:
                    speakers_map[text_before_colon] += " "

    return speakers_map


def build_corpus_csv(zip_file_path, output_path):

    data_list = []
    # open the Zip file
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        # extract Word files
        zip_ref.extractall('extracted_docs')
    # read each word file and extract the protocol entry number and type
    for filename in os.listdir('extracted_docs'):
        if filename.endswith(".docx"):
            # extract the wanted info from the file name
            file_path = os.path.join('extracted_docs', filename)
            entry_number = filename.split('_')[0]
            protocol_suf = filename.split('_')[1]
            protocol_type = "committee" if protocol_suf == "ptv" else "plenary"

            document = Document(file_path)
            speakers_file_map = extract_names_and_sentences(document, protocol_type)

            for speaker_name, speaker_text in speakers_file_map.items():

                sentences = split_sentence(speaker_text)
                for sentence in sentences:
                    cleaned_sentence = clean_sentence(sentence)
                    if cleaned_sentence == "":
                        continue
                    sentence_tokenize = tokenized(cleaned_sentence)
                    if calculate_sentence_length(sentence_tokenize) >= 4:
                        final_sentence = ' '.join(sentence_tokenize)

                        # adding the info to our data list
                        data_list.append({
                            'protocol_name': filename,
                            'number_knesset': entry_number,
                            'type_protocol': protocol_type,
                            'Speaker_name': speaker_name,
                            'Sentence_text': final_sentence
                        })

    df = pd.DataFrame(data_list)

    # the csv file path
    csv_file_path = output_path

    # with the data from the df to the csv file
    df.to_csv(csv_file_path, sep=',', encoding='utf-8', index=False)
    return csv_file_path


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit(1)

    zip_file_path = sys.argv[1]
    output_path = sys.argv[2]

    csv_path = build_corpus_csv(zip_file_path, output_path)
